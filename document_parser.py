"""
document_parser.py — Document Parser
CyberHealth AI | MedGemma Clinical Safety Auditor

Extracts text from PDF, Word (.docx), and Excel (.xlsx/.csv) files.
Accepts base64-encoded file content so it works over MCP protocol.
"""

import base64
import io
import logging

logger = logging.getLogger(__name__)


def extract_text_from_base64(file_base64: str, file_type: str) -> str:
    """
    Extracts plain text from a base64-encoded file.

    Args:
        file_base64: Base64-encoded file content
        file_type: One of 'pdf', 'docx', 'xlsx', 'csv'

    Returns:
        Extracted plain text string
    """
    try:
        file_bytes = base64.b64decode(file_base64)
        file_obj = io.BytesIO(file_bytes)
        file_type = file_type.lower().strip().lstrip(".")

        if file_type == "pdf":
            return _extract_pdf(file_obj)
        elif file_type == "docx":
            return _extract_docx(file_obj)
        elif file_type in ["xlsx", "xls"]:
            return _extract_excel(file_obj)
        elif file_type == "csv":
            return _extract_csv(file_obj)
        else:
            return f"Unsupported file type: {file_type}. Supported: pdf, docx, xlsx, csv"

    except Exception as e:
        logger.error(f"Document extraction failed: {e}")
        return f"Document extraction failed: {str(e)}"


def _extract_pdf(file_obj: io.BytesIO) -> str:
    """Extract text from PDF."""
    try:
        import PyPDF2
        reader = PyPDF2.PdfReader(file_obj)
        pages = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text and text.strip():
                pages.append(f"[Page {i+1}]\n{text.strip()}")
        if not pages:
            return "PDF appears to be scanned/image-based. No extractable text found."
        return "\n\n".join(pages)
    except ImportError:
        return "PyPDF2 not installed. Run: pip install PyPDF2"
    except Exception as e:
        return f"PDF extraction error: {str(e)}"


def _extract_docx(file_obj: io.BytesIO) -> str:
    """Extract text from Word document."""
    try:
        from docx import Document
        doc = Document(file_obj)
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text.strip())
        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells
                    if cell.text.strip()
                )
                if row_text:
                    paragraphs.append(row_text)
        if not paragraphs:
            return "Word document appears empty."
        return "\n".join(paragraphs)
    except ImportError:
        return "python-docx not installed. Run: pip install python-docx"
    except Exception as e:
        return f"Word extraction error: {str(e)}"


def _extract_excel(file_obj: io.BytesIO) -> str:
    """Extract text from Excel file."""
    try:
        import pandas as pd
        xl = pd.ExcelFile(file_obj)
        sheets = []
        for sheet_name in xl.sheet_names:
            df = xl.parse(sheet_name)
            df = df.dropna(how='all').fillna('')
            sheet_text = f"[Sheet: {sheet_name}]\n"
            sheet_text += df.to_string(index=False)
            sheets.append(sheet_text)
        if not sheets:
            return "Excel file appears empty."
        return "\n\n".join(sheets)
    except ImportError:
        return "pandas/openpyxl not installed. Run: pip install pandas openpyxl"
    except Exception as e:
        return f"Excel extraction error: {str(e)}"


def _extract_csv(file_obj: io.BytesIO) -> str:
    """Extract text from CSV file."""
    try:
        import pandas as pd
        df = pd.read_csv(file_obj)
        df = df.dropna(how='all').fillna('')
        return df.to_string(index=False)
    except ImportError:
        return "pandas not installed. Run: pip install pandas"
    except Exception as e:
        return f"CSV extraction error: {str(e)}"
